#!/usr/bin/env python3
"""
Extrae contenido de archivos DOCX a texto plano con formato Markdown.
Preserva tablas, párrafos, listas e imágenes (referencias).
"""

import sys
import json
from docx import Document
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from pathlib import Path


def extract_docx(file_path: str) -> dict:
    """
    Extrae el contenido de un archivo DOCX.

    Retorna:
        {
            'title': str,
            'paragraphs': list[str],
            'tables': list[dict],
            'has_images': bool,
            'error': str or None
        }
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        return {
            'title': Path(file_path).stem,
            'paragraphs': [],
            'tables': [],
            'has_images': False,
            'error': f'Error al abrir DOCX: {str(e)}'
        }

    # Título: intenta extraer del primer párrafo de estilo Heading 1
    title = Path(file_path).stem
    for para in doc.paragraphs:
        if para.style.name and 'Heading 1' in para.style.name and para.text.strip():
            title = para.text.strip()
            break

    # Extracto de párrafos y tablas manteniendo el orden
    content_blocks = []
    has_images = False

    for element in doc.element.body:
        # Párrafo
        if isinstance(element, CT_P):
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para and para.text.strip():
                # Preservar indentación si está presente
                text = para.text
                indent = ""
                if para.paragraph_format.left_indent:
                    indent = "  "
                content_blocks.append(indent + text)

            # Detectar imágenes en párrafos
            if para:
                for run in para.runs:
                    if run.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                        has_images = True

        # Tabla
        elif element.tag.endswith('}tbl'):
            # Encuentra la tabla en doc.tables
            for table in doc.tables:
                if table._element is element:
                    table_dict = {
                        'rows': [],
                        'cols': len(table.rows[0].cells) if table.rows else 0
                    }
                    for row in table.rows:
                        row_cells = []
                        for cell in row.cells:
                            cell_text = '\n'.join([p.text for p in cell.paragraphs]).strip()
                            row_cells.append(cell_text)
                        table_dict['rows'].append(row_cells)
                    content_blocks.append(('TABLE', table_dict))
                    break

    return {
        'title': title,
        'paragraphs': content_blocks,
        'has_images': has_images,
        'error': None
    }


def format_for_markdown(extracted: dict) -> str:
    """
    Formatea el contenido extraído como Markdown legible.
    """
    output = []

    if extracted['error']:
        return f"# Error\n\n{extracted['error']}"

    output.append(f"# {extracted['title']}\n")

    if extracted['has_images']:
        output.append("*Nota: Este documento contiene imágenes que no se pueden representar en texto plano.*\n")

    for block in extracted['paragraphs']:
        if isinstance(block, tuple) and block[0] == 'TABLE':
            table_dict = block[1]
            # Formato de tabla Markdown simple
            rows = table_dict['rows']
            if rows:
                # Header
                output.append('| ' + ' | '.join(rows[0]) + ' |')
                output.append('|' + '---|' * table_dict['cols'])
                # Body
                for row in rows[1:]:
                    output.append('| ' + ' | '.join(row) + ' |')
                output.append('')
        else:
            output.append(block)

    return '\n'.join(output)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(json.dumps({
            'error': 'Uso: extract_docx.py <ruta-docx>',
            'paragraphs': [],
            'title': '',
            'has_images': False
        }))
        sys.exit(1)

    file_path = sys.argv[1]
    result = extract_docx(file_path)

    if '--json' in sys.argv:
        print(json.dumps(result))
    else:
        print(format_for_markdown(result))
